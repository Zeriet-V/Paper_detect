from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ALIGN_MAP = {
    WD_PARAGRAPH_ALIGNMENT.LEFT: 'LEFT',
    WD_PARAGRAPH_ALIGNMENT.CENTER: 'CENTER',
    WD_PARAGRAPH_ALIGNMENT.RIGHT: 'RIGHT',
    WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'JUSTIFY',
    WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: 'DISTRIBUTE'
}

def fmt_align(a):
    if a is None: return None
    return ALIGN_MAP.get(a, str(a))

def inspect_alignments(docx_path):
    doc = Document(docx_path)
    print(f"打开: {docx_path}，段落数: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs, 1):
        direct = p.paragraph_format.alignment  # 直接格式
        style_align = None
        try:
            style_align = p.style.paragraph_format.alignment  # 样式定义里的对齐
        except Exception:
            style_align = None
        effective = p.alignment  # python-docx 有时返回计算后的 alignment（或 None）
        print(f"段落 {i}: \"{p.text[:50].strip()}\"")
        print(f"  直接(paragraph_format.alignment): {fmt_align(direct)}")
        print(f"  样式(style.paragraph_format.alignment): {fmt_align(style_align)}")
        print(f"  paragraph.alignment (可能为计算值): {fmt_align(effective)}\n")


inspect_alignments('D:\pyproject\paper_detect\\template\\test_abstract.docx')
inspect_alignments('D:\pyproject\paper_detect\\template\\test.docx')
