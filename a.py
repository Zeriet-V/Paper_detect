from docx import Document
from docx.shared import Inches
import docx

# 版本校验
print(f"当前版本: {docx.__version__}")  # 1.2.0

# 创建文档与段落
doc = Document()
para = doc.add_paragraph("这是一个示例段落，用于演示批注功能。")
run = para.runs[0]

# 1. 添加基本批注
comment1 = doc.add_comment(runs=run, text="这是一个简单批注", author="张三", initials="ZS")
print("已添加基本批注")

# 2. 读取并修改批注
for comment in doc.comments:
    print(f"批注文本: {comment.text}, 作者: {comment.author}, 缩写: {comment.initials}")
    comment.author = "李四"
    comment.initials = "LS"
print("已修改批注作者和缩写")

# 3. 添加格式化文本批注
comment2 = doc.add_comment(runs=run, text="", author="王五", initials="WW")
cmt_para = comment2.add_paragraph()
cmt_run1 = cmt_para.add_run("这是一个")
cmt_run1.bold = True
cmt_run2 = cmt_para.add_run("丰富文本批注")
cmt_run2.italic = True
print("已添加带丰富文本的批注")

# 4. 添加表格批注
comment3 = doc.add_comment(runs=run, text="", author="赵六", initials="ZL")
table = comment3.add_table(rows=2, cols=2, width=Inches(6))
table.style = 'Table Grid'
table.cell(0, 0).text = '姓名'
table.cell(0, 1).text = '年龄'
table.cell(1, 0).text = '张三'
table.cell(1, 1).text = '30'
print("已添加带表格的批注")

# 保存
doc.save("comment_demo.docx")
print("文档已保存为 comment_demo.docx")
