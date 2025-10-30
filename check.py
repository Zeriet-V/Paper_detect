#!/usr/bin/env python3
"""
读取 docx 中每个段落的格式信息并打印出来（字体、字号、对齐方式、缩进、段前/段后、行距等）。
依赖：python-docx（pip install python-docx）

用法：
    python read_paragraph_format.py path/to/file.docx

说明：
- Word 的格式可能来自“直接格式设置（direct formatting）”或来自样式（style）。脚本会尝试显示两者。
- 有些属性在 docx 中为 None（表示未设置），脚本会以 None 显示。
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ALIGN_MAP = {
    WD_PARAGRAPH_ALIGNMENT.LEFT: 'LEFT',
    WD_PARAGRAPH_ALIGNMENT.CENTER: 'CENTER',
    WD_PARAGRAPH_ALIGNMENT.RIGHT: 'RIGHT',
    WD_PARAGRAPH_ALIGNMENT.JUSTIFY: 'JUSTIFY',
    WD_PARAGRAPH_ALIGNMENT.DISTRIBUTE: 'DISTRIBUTE'
}


def to_pts(length_obj):
    """把 docx length 对象转换为点值（pt）。如果为 None，返回 None。"""
    if length_obj is None:
        return None
    # length_obj 通常有 .pt 属性，但有时不是长度对象，捕获异常
    try:
        return float(length_obj.pt)
    except Exception:
        return str(length_obj)


def get_effective_alignment(paragraph):
    """尝试返回段落最终生效的对齐字符串（优先直接设置，再看样式）。"""
    if paragraph.alignment is not None:
        return ALIGN_MAP.get(paragraph.alignment, str(paragraph.alignment))
    # 尝试从样式中取
    try:
        sty_align = paragraph.style.paragraph_format.alignment
        if sty_align is not None:
            return ALIGN_MAP.get(sty_align, str(sty_align))
    except Exception:
        pass
    return None


def summarize_paragraph(paragraph, index):
    text_preview = paragraph.text.strip().replace('\n', ' ')[:200]
    style_name = getattr(paragraph.style, 'name', None)

    # 段落格式
    pf = paragraph.paragraph_format
    left_indent = to_pts(getattr(pf, 'left_indent', None))
    right_indent = to_pts(getattr(pf, 'right_indent', None))
    first_line_indent = to_pts(getattr(pf, 'first_line_indent', None))
    space_before = to_pts(getattr(pf, 'space_before', None))
    space_after = to_pts(getattr(pf, 'space_after', None))
    line_spacing = getattr(pf, 'line_spacing', None)
    if line_spacing is None:
        line_spacing_pt = None
    else:
        # 有时 line_spacing 是 float（倍行距），有时是 Length（有 .pt）
        try:
            line_spacing_pt = float(line_spacing.pt)
        except Exception:
            line_spacing_pt = line_spacing

    alignment = get_effective_alignment(paragraph)

    # run 层的字体信息（可能每个 run 不同）
    run_fonts = set()
    run_sizes = set()
    any_bold = False
    any_italic = False
    for run in paragraph.runs:
        font = run.font
        if font is None:
            continue
        if font.name:
            run_fonts.add(font.name)
        if font.size is not None:
            try:
                run_sizes.add(float(font.size.pt))
            except Exception:
                run_sizes.add(str(font.size))
        if run.bold:
            any_bold = True
        if run.italic:
            any_italic = True

    # 来自样式的字体信息（如果有）
    style_font_name = None
    style_font_size_pt = None
    try:
        sty_font = paragraph.style.font
        style_font_name = getattr(sty_font, 'name', None)
        sty_size = getattr(sty_font, 'size', None)
        if sty_size is not None:
            try:
                style_font_size_pt = float(sty_size.pt)
            except Exception:
                style_font_size_pt = sty_size
    except Exception:
        pass

    summary = {
        'index': index,
        'text_preview': text_preview,
        'style_name': style_name,
        'alignment': alignment,
        'left_indent_pt': left_indent,
        'right_indent_pt': right_indent,
        'first_line_indent_pt': first_line_indent,
        'space_before_pt': space_before,
        'space_after_pt': space_after,
        'line_spacing_pt_or_value': line_spacing_pt,
        'run_fonts': sorted(list(run_fonts)) if run_fonts else None,
        'run_font_sizes_pt': sorted(list(run_sizes)) if run_sizes else None,
        'any_run_bold': any_bold,
        'any_run_italic': any_italic,
        'style_font_name': style_font_name,
        'style_font_size_pt': style_font_size_pt
    }
    return summary


def print_summary(s):
    print(f"段落 {s['index']}: \"{s['text_preview']}\"")
    print(f"  样式: {s['style_name']}")
    print(f"  对齐: {s['alignment']}")
    print(f"  缩进 - 左: {s['left_indent_pt']} pt, 右: {s['right_indent_pt']} pt, 首行: {s['first_line_indent_pt']} pt")
    print(f"  段距 - 段前: {s['space_before_pt']} pt, 段后: {s['space_after_pt']} pt")
    print(f"  行距(或值): {s['line_spacing_pt_or_value']}")
    print(f"  run 层字体: {s['run_fonts']}, run 字号(pt): {s['run_font_sizes_pt']}")
    print(f"  run 中有加粗: {s['any_run_bold']}, 斜体: {s['any_run_italic']}")
    print(f"  来自样式的字体: {s['style_font_name']}, 样式字号(pt): {s['style_font_size_pt']}\n")


def main(docx_path):
    doc = Document(docx_path)
    print(f"打开: {docx_path}, 段落数: {len(doc.paragraphs)}\n")
    for i, p in enumerate(doc.paragraphs, start=1):
        s = summarize_paragraph(p, i)
        print_summary(s)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python read_paragraph_format.py path/to/file.docx')
        sys.exit(1)
    path = sys.argv[1]
    main(path)
